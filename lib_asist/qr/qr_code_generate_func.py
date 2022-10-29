import qrcode
from PIL import Image
import os
from docx import Document
from docx.shared import Inches

def generate_qr_cods(n,m):
    """
    qr cod ni generate qiladi
    :param n: boshlangich id
    :param m: generate qilinadigan id larning minimal soni
    :return: [ohirgi id,word fayl ning adresi]
    """
    document = Document()
    background = Image.new(mode="RGB", size=(1900, 370),
                   color=(10000, 1000, 1000))
    wh = [(0,0),(380,0),(760,0),(1140,0),(1520,0)]
    whn,t = 0,False
    for i in range(int(n),int(n+m)):
        data = 'https://ubtuit-library.uz/qr/'+str(i)
        frontImage = qrcode.make(data)
        frontImage = frontImage.convert("RGBA")
        background.paste(frontImage, wh[whn], frontImage)
        if whn==4:
            t=False
            whn=0
            background.save('test.jpg')
            document.add_picture('test.jpg',width=Inches(6))
            os.remove('test.jpg')
            background = Image.new(mode="RGB", size=(1900, 370),
                                   color=(10000, 1000, 1000))
        else:
            t=True
            whn+=1
    if t:
        background.save('test.jpg')
        document.add_picture('test.jpg', width=Inches(6))
        os.remove('test.jpg')
    return [document,(n+m)]
a=generate_qr_cods(20,11)
a[0].save('doc.docx')